"""
文档加载器 — 从 data/raw 目录加载法律文本文件（支持 txt 和 docx）
含日期元数据提取，支持时效感知
"""
import csv
import os
import json
import re
from typing import List, Dict, Tuple, Optional
from datetime import datetime
from docx import Document as DocxDocument
from langchain_core.documents import Document
from src.config import Config


def _read_file(filepath: str) -> str:
    """读取 txt 或 docx 文件，返回文本内容"""
    if filepath.endswith(".docx"):
        doc = DocxDocument(filepath)
        return "\n".join(p.text for p in doc.paragraphs)
    else:
        with open(filepath, "r", encoding="utf-8") as f:
            return f.read()


def _strip_ext(filename: str) -> str:
    """去掉文件扩展名"""
    for ext in [".txt", ".docx"]:
        if filename.endswith(ext):
            return filename[: -len(ext)]
    return filename


def _parse_date_from_filename(filename: str) -> Optional[str]:
    """从文件名末尾提取 _YYYYMMDD 格式的日期，返回 YYYY-MM-DD"""
    m = re.search(r'_(\d{4})(\d{2})(\d{2})(?:\.(?:txt|docx))?$', filename)
    if m:
        return f"{m.group(1)}-{m.group(2)}-{m.group(3)}"
    return None


def _parse_effective_date_from_content(text: str) -> Optional[str]:
    """从法律文本内容中提取施行日期（生效日期）"""
    # 模式1: 自XXXX年XX月XX日起施行
    m = re.search(r'自(\d{4})年(\d{1,2})月(\d{1,2})日起施行', text)
    if m:
        return f"{m.group(1)}-{int(m.group(2)):02d}-{int(m.group(3)):02d}"

    # 模式2: XXXX年XX月XX日施行 / XXXX年XX月XX日起施行
    m = re.search(r'(\d{4})年(\d{1,2})月(\d{1,2})日.*?施行', text)
    if m:
        return f"{m.group(1)}-{int(m.group(2)):02d}-{int(m.group(3)):02d}"

    # 模式3: 发布日期（用于司法解释）— "XXXX年XX月XX日最高人民法院"
    m = re.search(r'(\d{4})年(\d{1,2})月(\d{1,2})日', text[:500])
    if m:
        return f"{m.group(1)}-{int(m.group(2)):02d}-{int(m.group(3)):02d}"

    return None


def _parse_publish_date_from_content(text: str) -> Optional[str]:
    """提取发布日期"""
    m = re.search(r'发布时间[：:]?\s*(\d{4})-(\d{1,2})-(\d{1,2})', text)
    if m:
        return f"{m.group(1)}-{int(m.group(2)):02d}-{int(m.group(3)):02d}"
    return None


def _extract_law_name_base(filename: str) -> str:
    """从文件名提取法律基础名称（去掉日期后缀和扩展名）
    如 '中华人民共和国劳动法_20181229.docx' -> '中华人民共和国劳动法'
    """
    name = _strip_ext(filename)
    # 去掉 _YYYYMMDD 后缀
    name = re.sub(r'_\d{8}$', '', name)
    return name


# ── 条文切分辅助 ─────────────────────────────────

# 只匹配主条文标题：条号后须跟全角空格/空白/左括号，避免命中正文内的条文引用（如"依照本法第四十六条规定"）
ARTICLE_RE = re.compile(r'第([一二三四五六七八九十百千零]+)条(?=[　\s（])')
_CN_DIGITS = {'零': 0, '一': 1, '二': 2, '三': 3, '四': 4, '五': 5, '六': 6, '七': 7, '八': 8, '九': 9}


def cn_num_to_int(s: str) -> Optional[int]:
    """中文数字转整数：四十六 -> 46；一百零二 -> 102"""
    if not s:
        return None
    total, section = 0, 0
    for ch in s:
        if ch in _CN_DIGITS:
            section = _CN_DIGITS[ch]
        elif ch == '十':
            total += (section or 1) * 10
            section = 0
        elif ch == '百':
            total += (section or 1) * 100
            section = 0
        elif ch == '千':
            total += (section or 1) * 1000
            section = 0
        else:
            return None
    return total + section


def split_articles(text: str) -> List[Tuple[str, str]]:
    """将法律全文切分为 [(条文号如'第四十六条', 条文正文), ...]"""
    matches = list(ARTICLE_RE.finditer(text))
    articles = []
    for i, m in enumerate(matches):
        end = matches[i + 1].start() if i + 1 < len(matches) else len(text)
        articles.append((m.group(0), text[m.end():end].strip()))
    return articles


def first_article_no(text: str) -> str:
    """提取文本中出现的第一个条文号（如'第四十六条'），无则空串"""
    m = ARTICLE_RE.search(text)
    return m.group(0) if m else ""


class LegalDocumentLoader:
    """加载法律法规、司法解释和案例文本（含日期元数据）"""

    def __init__(self, data_dir: str = None):
        self.data_dir = data_dir or Config.DATA_DIR

    # ── 日期提取辅助 ─────────────────────────────────

    def _extract_dates(self, filepath: str, content: str) -> Dict[str, str]:
        """从文件名和内容综合提取日期信息
        Returns: {"publish_date": "...", "effective_date": "..."}
        """
        filename = os.path.basename(filepath)
        dates: Dict[str, str] = {}

        # 从文件名提取
        fn_date = _parse_date_from_filename(filename)
        if fn_date:
            dates["publish_date"] = fn_date
            dates["effective_date"] = fn_date

        # 从内容提取（覆盖文件名日期）
        eff_date = _parse_effective_date_from_content(content)
        if eff_date:
            dates["effective_date"] = eff_date

        pub_date = _parse_publish_date_from_content(content)
        if pub_date:
            dates["publish_date"] = pub_date

        # 如果只有 effective_date 没有 publish_date，用 effective 代替
        if "effective_date" in dates and "publish_date" not in dates:
            dates["publish_date"] = dates["effective_date"]
        if "publish_date" in dates and "effective_date" not in dates:
            dates["effective_date"] = dates["publish_date"]

        return dates

    @staticmethod
    def _infer_status(effective_date: Optional[str], is_latest_version: bool) -> str:
        """推断法律效力状态"""
        if not effective_date:
            return "效力未知"
        if not is_latest_version:
            return "已被修订"
        try:
            eff_dt = datetime.strptime(effective_date, "%Y-%m-%d")
            if eff_dt > datetime.now():
                return "尚未生效"
        except ValueError:
            pass
        return "现行有效"

    # ── 版本检测 ─────────────────────────────────────

    def _find_latest_versions(self, statutes_dir: str) -> Dict[str, str]:
        """扫描 statutes 目录，找出每部法律的最新版本文件名
        Returns: {法律基础名: 最新版本文件名}
        """
        if not os.path.exists(statutes_dir):
            return {}

        law_versions: Dict[str, List[Tuple[str, Optional[str]]]] = {}

        for filename in os.listdir(statutes_dir):
            if not (filename.endswith(".txt") or filename.endswith(".docx")):
                continue
            if os.path.getsize(os.path.join(statutes_dir, filename)) < 100:
                continue  # 跳过空存根

            base = _extract_law_name_base(filename)
            date_str = _parse_date_from_filename(filename)
            if base not in law_versions:
                law_versions[base] = []
            law_versions[base].append((filename, date_str))

        latest: Dict[str, str] = {}
        for base, versions in law_versions.items():
            # 按日期排序，取最新
            dated = [(f, d) for f, d in versions if d]
            if dated:
                dated.sort(key=lambda x: x[1], reverse=True)
                latest[base] = dated[0][0]
            elif versions:
                latest[base] = versions[0][0]

        return latest

    # ── 文档加载 ─────────────────────────────────────

    def load_all(self) -> List[Document]:
        """加载所有文档"""
        docs = []
        docs.extend(self.load_statutes())
        docs.extend(self.load_interpretations())
        docs.extend(self.load_cases_text())
        return docs

    def load_statutes(self) -> List[Document]:
        """加载法律法规文件（txt 或 docx），含日期元数据"""
        docs = []
        statutes_dir = os.path.join(self.data_dir, "statutes")
        if not os.path.exists(statutes_dir):
            return docs

        latest_map = self._find_latest_versions(statutes_dir)

        for filename in os.listdir(statutes_dir):
            if not (filename.endswith(".txt") or filename.endswith(".docx")):
                continue
            filepath = os.path.join(statutes_dir, filename)

            # 跳过空存根（< 100 字节）
            if os.path.getsize(filepath) < 100:
                continue

            content = _read_file(filepath)
            base_name = _extract_law_name_base(filename)
            law_name = base_name

            # 提取日期
            dates = self._extract_dates(filepath, content)

            # 判断是否最新版本
            is_latest = (latest_map.get(base_name, filename) == filename)
            status = self._infer_status(dates.get("effective_date"), is_latest)

            docs.append(
                Document(
                    page_content=content,
                    metadata={
                        "source": filepath,
                        "doc_type": "statute",
                        "law_name": law_name,
                        "publish_date": dates.get("publish_date", ""),
                        "effective_date": dates.get("effective_date", ""),
                        "expiry_date": "",
                        "status": status,
                        "is_latest": is_latest,
                    },
                )
            )
        return docs

    def load_interpretations(self) -> List[Document]:
        """加载司法解释文件（txt 或 docx），含日期元数据"""
        docs = []
        interp_dir = os.path.join(self.data_dir, "interpretations")
        if not os.path.exists(interp_dir):
            return docs

        for filename in os.listdir(interp_dir):
            if not (filename.endswith(".txt") or filename.endswith(".docx")):
                continue
            filepath = os.path.join(interp_dir, filename)

            if os.path.getsize(filepath) < 100:
                continue

            content = _read_file(filepath)
            dates = self._extract_dates(filepath, content)
            title = _strip_ext(filename)

            docs.append(
                Document(
                    page_content=content,
                    metadata={
                        "source": filepath,
                        "doc_type": "interpretation",
                        "title": title,
                        "publish_date": dates.get("publish_date", ""),
                        "effective_date": dates.get("effective_date", ""),
                        "expiry_date": "",
                        "status": self._infer_status(dates.get("effective_date"), True),
                    },
                )
            )
        return docs

    def load_cases_text(self) -> List[Document]:
        """加载案例文件（JSON 或 CSV），转为可检索文本"""
        docs = []
        cases_dir = os.path.join(self.data_dir, "cases")
        if not os.path.exists(cases_dir):
            return docs

        for filename in os.listdir(cases_dir):
            filepath = os.path.join(cases_dir, filename)

            if filename.endswith(".csv"):
                cases = self._read_cases_from_csv(filepath)
            elif filename.endswith(".json"):
                cases = self._read_cases_from_json(filepath)
            else:
                continue

            for case in cases:
                text_parts = []

                def _str(val):
                    if isinstance(val, list):
                        return '、'.join(val)
                    return str(val) if val else ''

                if case.get("case_content"):
                    text_parts.append(f"案例内容：{_str(case['case_content'])}")
                if case.get("issues"):
                    text_parts.append(f"争议焦点：{_str(case['issues'])}")
                if case.get("reasoning"):
                    text_parts.append(f"法院认为：{_str(case['reasoning'])}")
                if case.get("judgment"):
                    text_parts.append(f"判决结果：{_str(case['judgment'])}")
                if case.get("keywords"):
                    text_parts.append(f"关键词：{_str(case['keywords'])}")

                docs.append(
                    Document(
                        page_content="\n".join(text_parts),
                        metadata={
                            "source": filepath,
                            "doc_type": "case",
                            "case_number": case.get("case_number", ""),
                            "court": case.get("court", ""),
                            "judge_date": case.get("judge_date", ""),
                            "keywords": case.get("keywords", "") if isinstance(case.get("keywords"), list)
                                        else _str(case.get("keywords", "")).split(";"),
                        },
                    )
                )
        return docs

    @staticmethod
    def _read_cases_from_csv(filepath: str) -> List[Dict]:
        cases = []
        with open(filepath, "r", encoding="utf-8-sig") as f:
            reader = csv.DictReader(f)
            for row in reader:
                cases.append({
                    "case_number": row.get("case_number", "").strip(),
                    "court": row.get("court", "").strip(),
                    "judge_date": row.get("judge_date", "").strip(),
                    "case_content": row.get("case_content", "").strip(),
                    "issues": row.get("issues", "").strip(),
                    "reasoning": row.get("reasoning", "").strip(),
                    "judgment": row.get("judgment", "").strip(),
                    "legal_basis": row.get("legal_basis", "").strip(),
                    "keywords": row.get("keywords", "").strip(),
                })
        return cases

    @staticmethod
    def _read_cases_from_json(filepath: str) -> List[Dict]:
        with open(filepath, "r", encoding="utf-8") as f:
            data = json.load(f)
        if isinstance(data, list):
            return data
        return [data]

    def load_cases_json(self) -> List[Dict]:
        """加载案例原始数据，用于导入 Neo4j"""
        cases = []
        cases_dir = os.path.join(self.data_dir, "cases")
        if not os.path.exists(cases_dir):
            return cases

        for filename in os.listdir(cases_dir):
            filepath = os.path.join(cases_dir, filename)
            if filename.endswith(".csv"):
                cases.extend(self._read_cases_from_csv(filepath))
            elif filename.endswith(".json"):
                cases.extend(self._read_cases_from_json(filepath))
        return cases
